from pathlib import Path

import pytest
from docx import Document

from kami_excel_extractor.docx_renderer import DocxRenderer


def test_docx_rendering_hybrid_structure(tmp_path):
    # ダミー画像を生成
    from PIL import Image
    media_dir = tmp_path / "media"
    media_dir.mkdir(parents=True, exist_ok=True)
    dummy_image = media_dir / "chart.png"
    img = Image.new("RGB", (100, 100), color="red")
    img.save(dummy_image)

    renderer = DocxRenderer(tmp_path)

    # 構造化データ (実データ同様に raw_media情報を表す data キーと、論理情報の photos キーの両方を含む)
    structured_data = {
        "sheets": {
            "現場写真": {
                "summary": "現場写真のリストです。",
                "data": [
                    {"coord": "A3", "filename": "chart.png", "type": "image"}
                ],
                "photos": [
                    {
                        "photo_no": "No.1",
                        "title": "クラック発生状況",
                        "coordinate_title": "A1",
                        "coordinate_image": "A3",
                        "image_details": "赤い太線で囲まれたクラック近接画像"
                    }
                ]
            }
        }
    }

    # 原本データ
    raw_data = {
        "sheets": {
            "現場写真": {
                "cells": [
                    {"coord": "A1", "row": 1, "col": 1, "value": "写真 No.1: クラック発生状況", "style": {"borders": {}, "bold": True}}
                ],
                "media": [
                    {"coord": "A3", "filename": "chart.png", "visual_summary": "画像データ"}
                ]
            }
        }
    }

    docx_path = renderer.generate_docx(
        structured_data=structured_data,
        raw_data=raw_data,
        source_filename="hybrid_report.xlsx",
        output_name="hybrid_output"
    )

    assert docx_path.exists()
    doc = Document(docx_path)

    # 1. 重複テーブル排除の検証
    # photosキーがあるため、dataキー（生ダンプ表）はスキップされ、論理テーブルの1つのみが出力されていること
    assert len(doc.tables) == 1
    table = doc.tables[0]

    headers = [cell.text for cell in table.rows[0].cells]
    assert "photo_no" in headers
    assert "title" in headers
    # dataキーのヘッダーが含まれていないこと
    assert "filename" not in headers

    # テーブル内に画像オブジェクトが含まれていないこと
    table_xml = "".join(cell._tc.xml for row in table.rows for cell in row.cells)
    assert "w:drawing" not in table_xml

    # 2. テーブル外の画像とキャプションの関連付け検証
    captions = [p.text for p in doc.paragraphs if p.style.name == "Caption"]
    assert any("No.1" in c and "表中座標: A1" in c and "画像座標: A3" in c for c in captions)

    # 段落（テーブル外）の中に画像 (w:drawing) が含まれていること
    all_paragraphs_xml = "".join(p._p.xml for p in doc.paragraphs)
    assert "w:drawing" in all_paragraphs_xml

    # 3. _raw_data などのデバッグ用メタデータがDOCXに出力されていないこと
    all_text = "\n".join(p.text for p in doc.paragraphs)
    assert "_raw_data" not in all_text


def test_docx_fallback_to_raw_table(tmp_path):
    renderer = DocxRenderer(tmp_path)

    # 構造化データが完全に空 (LLM構造化なし)
    structured_data = {
        "sheets": {
            "報告概要": {}
        }
    }

    # 原本データがある
    raw_data = {
        "sheets": {
            "報告概要": {
                "cells": [
                    {"coord": "A1", "row": 1, "col": 1, "value": "生セルデータ", "colspan": 1, "rowspan": 1, "style": {"borders": {}, "bold": False}}
                ],
                "media": []
            }
        }
    }

    docx_path = renderer.generate_docx(
        structured_data=structured_data,
        raw_data=raw_data,
        source_filename="fallback_report.xlsx",
        output_name="fallback_output"
    )

    assert docx_path.exists()
    doc = Document(docx_path)

    # 構造化データがないため、フォールバックとして原本テーブルが出力されていること
    assert len(doc.tables) == 1
    table = doc.tables[0]

    # セルの値にプレフィックス付きで生セルデータが入っていること
    assert "[A1] 生セルデータ" in table.cell(0, 0).text
