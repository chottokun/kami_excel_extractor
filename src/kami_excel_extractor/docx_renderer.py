import logging
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Inches

logger = logging.getLogger(__name__)

def col_to_num(col_str: str) -> int:
    """Excelの列文字列(A, B, C...)をインデックス(1-indexed)に変換する。"""
    num = 0
    for c in col_str.upper():
        num = num * 26 + (ord(c) - ord('A') + 1)
    return num

def parse_coord(coord_str: str) -> Tuple[int, int]:
    """セル座標文字列(A1, B5...)を (row, col) インデックス(1-indexed)に変換する。"""
    match = re.match(r"^([A-Z]+)([0-9]+)$", coord_str.upper())
    if not match:
        return (0, 0)
    col_str, row_str = match.groups()
    return int(row_str), col_to_num(col_str)

def is_coord_in_range(coord: str, start: str, end: str) -> bool:
    """座標 coord が start から end の範囲内に含まれるか判定する。"""
    r_img, c_img = parse_coord(coord)
    r_start, c_start = parse_coord(start)
    r_end, c_end = parse_coord(end)
    if r_img == 0 or r_start == 0 or r_end == 0:
        return False
    return (r_start <= r_img <= r_end) and (c_start <= c_img <= c_end)


class DocxRenderer:
    """
    LLMが抽出した意味的・構造化データをベースに、
    Difyナレッジベースに最適な論理的ドキュメントとしてDOCXファイルを再構成するクラス。
    ハイブリッド構造（テーブル内にインサイトテキスト、テーブル外に関連明記画像とキャプション）を構築する。
    """

    def __init__(self, output_dir: Path):
        self.output_dir = Path(output_dir).resolve()
        self.media_dir = self.output_dir / "media"

    def generate_docx(
        self,
        structured_data: Dict[str, Any],
        raw_data: Dict[str, Any],
        source_filename: str,
        output_name: Optional[str] = None,
        include_logic_annotations: bool = True,
        image_width_inches: float = 5.0,
    ) -> Path:
        """
        Dify最適化DOCXファイルを生成する（ハイブリッド構造版）。
        """
        doc = Document()

        # 1. ドキュメントタイトル (Heading 1)
        doc.add_heading(source_filename, level=1)

        sheets_structured = structured_data.get("sheets", {})
        sheets_raw = raw_data.get("sheets", {})

        for sheet_name, sheet_data in sheets_structured.items():
            raw_sheet = sheets_raw.get(sheet_name, {})
            raw_cells = raw_sheet.get("cells", [])
            raw_media = raw_sheet.get("media", [])

            # シート見出し
            doc.add_heading(sheet_name, level=1)

            # シート全体の概要・説明
            self._add_overview_text(doc, sheet_data)

            # 挿入済み画像を追跡するセット
            inserted_media = set()

            # 構造化データ（論理テーブルや概要データ）があるか判定
            table_data, _ = self._detect_table_data(sheet_data)
            has_structured_data = table_data is not None or (isinstance(sheet_data.get("data"), dict) and len(sheet_data["data"]) > 0)

            if has_structured_data:
                # 構造化データが存在する場合は、論理データのみを出力
                self._render_structured_content(doc, sheet_data, raw_cells, raw_media, inserted_media)
            else:
                # 構造化データがない場合は、原本テーブルをフォールバック出力
                if raw_cells:
                    self._add_table_with_merges_and_prefixes(doc, raw_cells)

            # テーブル外（直後）に関連画像を挿入
            self._add_associated_images_below_table(doc, sheet_data, raw_media, inserted_media, image_width_inches)

            # 計算式情報の注釈
            if include_logic_annotations:
                self._add_logic_annotations(doc, raw_cells)

            # テーブルに組み込まれなかった残りの画像を出力（フォールバック）
            for media_item in raw_media:
                filename = media_item.get("filename")
                if filename and filename not in inserted_media:
                    self._add_image_with_caption(doc, media_item, image_width_inches)

            # シートごとのページ区切り
            doc.add_page_break()

        # 出力ファイル名の解決
        if not output_name:
            output_name = Path(source_filename).stem

        out_path = self.output_dir / f"{output_name}.docx"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        doc.save(out_path)
        logger.info(f"Generated Dify-optimized DOCX: {out_path}")
        return out_path

    def _add_overview_text(self, doc: Document, sheet_data: Dict[str, Any]) -> None:
        """
        シートの説明や概要を最上部に挿入する。
        """
        for key in ["summary", "description", "overview", "analysis"]:
            if key in sheet_data and sheet_data[key]:
                doc.add_heading("解析サマリー", level=2)
                doc.add_paragraph(str(sheet_data[key]))
                break

    def _find_cell_coord(self, raw_cells: List[Dict[str, Any]], value: Any, key_name: Optional[str] = None) -> Optional[str]:
        """
        値が一致する原本セルを検索し、その座標 (A1等) を返す。
        """
        if value is None:
            return None
        val_str = str(value).strip()
        if not val_str:
            return None

        # 辞書/リスト型の場合は内部のテキストを対象とする
        if isinstance(value, (dict, list)):
            val_str = str(value)

        # 完全一致を優先して探索
        for cell in raw_cells:
            cell_val = cell.get("value")
            if cell_val is not None and str(cell_val).strip() == val_str:
                return cell.get("coord")

        # 部分一致で探索
        for cell in raw_cells:
            cell_val = cell.get("value")
            if cell_val is not None and val_str in str(cell_val):
                return cell.get("coord")

        # キー名での探索
        if key_name:
            key_str = str(key_name).strip()
            for cell in raw_cells:
                cell_val = cell.get("value")
                if cell_val is not None and key_str in str(cell_val):
                    return cell.get("coord")

        return None

    def _detect_table_data(self, sheet_data: Dict[str, Any]) -> Tuple[Optional[List[Dict[str, Any]]], Optional[str]]:
        """
        sheet_data 内から最も意味論的に価値の高い論理テーブルデータを検出し、そのリストとキー名を返す。
        カスタム意味リスト(photos等)を data キー（生データリスト）より優先する。
        """
        # 1. photos や records などのカスタム意味的リストキーを最優先で探索
        for k, v in sheet_data.items():
            if k not in ["data", "summary", "description", "overview", "analysis", "_raw_data", "raw_data"]:
                if isinstance(v, list) and len(v) > 0 and isinstance(v[0], dict):
                    return v, k

        # 2. カスタムキーがなければ、一般的な "data" キーを探索
        if "data" in sheet_data and isinstance(sheet_data["data"], list):
            return sheet_data["data"], "data"

        return None, None

    def _render_structured_content(
        self,
        doc: Document,
        sheet_data: Dict[str, Any],
        raw_cells: List[Dict[str, Any]],
        raw_media: List[Dict[str, Any]],
        inserted_media: set
    ) -> None:
        """
        構造化データ (data) を論理テーブルまたは文章として再構築する。
        """
        table_data, table_key = self._detect_table_data(sheet_data)

        if table_data:
            # 検出された辞書リストからテーブルを作成
            self._add_structured_table_with_insights(doc, table_data, raw_cells, raw_media, inserted_media)

            # テーブルデータ以外の追加の意味データがあれば段落として展開
            # 重複を防ぐため、k != table_key に加え "data" も除外する
            ignored_keys = ["summary", "description", "overview", "analysis", "_raw_data", "raw_data", "data"]
            for k, v in sheet_data.items():
                if k != table_key and k not in ignored_keys:
                    if isinstance(v, (dict, list)):
                        doc.add_heading(str(k), level=3)
                        self._render_structured_content(doc, {"data": v}, raw_cells, raw_media, inserted_media)
                    else:
                        p = doc.add_paragraph()
                        p.add_run(f"{k}: ").bold = True
                        p.add_run(str(v))
        else:
            # プリミティブな値やネストした辞書データの通常の展開
            data = sheet_data.get("data")
            if data:
                if isinstance(data, dict):
                    for k, v in data.items():
                        if isinstance(v, (dict, list)):
                            doc.add_heading(str(k), level=3)
                            self._render_structured_content(doc, {"data": v}, raw_cells, raw_media, inserted_media)
                        else:
                            p = doc.add_paragraph()
                            coord = self._find_cell_coord(raw_cells, v, k)
                            prefix = f"[{coord}] " if coord else ""
                            run_k = p.add_run(f"{prefix}{k}: ")
                            run_k.bold = True
                            p.add_run(str(v))
                else:
                    doc.add_paragraph(str(data))

    def _add_table_with_merges_and_prefixes(self, doc: Document, cells: List[Dict[str, Any]]) -> None:
        """
        原本の結合セルレイアウトをWordテーブルで再現し、各セルに [A1] などの座標プレフィックスを付与する。
        (フォールバック用)
        """
        if not cells:
            return

        max_row = max(c["row"] for c in cells)
        max_col = max(c["col"] for c in cells)

        table = doc.add_table(rows=max_row, cols=max_col)
        table.style = 'Table Grid'

        # 1. データの充填
        for cell in cells:
            r_idx = cell["row"] - 1
            c_idx = cell["col"] - 1
            val = cell.get("value")
            val_str = str(val) if val is not None else ""
            coord = cell.get("coord", "")

            if val is not None and coord:
                cell_text = f"[{coord}] {val_str}"
            else:
                cell_text = val_str

            table.cell(r_idx, c_idx).text = cell_text

            is_header = (cell["row"] == 1) or cell.get("style", {}).get("bold", False)
            if is_header and cell_text:
                for paragraph in table.cell(r_idx, c_idx).paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # 2. 結合の適用
        for cell in cells:
            colspan = cell.get("colspan", 1)
            rowspan = cell.get("rowspan", 1)

            if colspan > 1 or rowspan > 1:
                r_start = cell["row"] - 1
                c_start = cell["col"] - 1
                r_end = r_start + rowspan - 1
                c_end = c_start + colspan - 1

                try:
                    cell_start = table.cell(r_start, c_start)
                    cell_end = table.cell(r_end, c_end)
                    if cell_start != cell_end:
                        cell_start.merge(cell_end)
                except Exception as e:
                    logger.warning(f"Failed to merge cells from ({r_start},{c_start}) to ({r_end},{c_end}): {e}")

    def _find_associated_media(self, row_data: Dict[str, Any], raw_media: List[Dict[str, Any]]) -> Optional[Dict[str, Any]]:
        """
        行データに関連する画像を raw_media から見つける。
        """
        # 1. coordinate_image, coordinate, coord キーでの直接比較
        coord = row_data.get("coordinate_image") or row_data.get("coordinate") or row_data.get("coord")
        if coord:
            coord_str = str(coord).upper()
            for m in raw_media:
                m_coord = str(m.get("coord", "")).upper()
                if m_coord == coord_str:
                    return m

        # 2. photo_area などの範囲オブジェクトでの比較
        photo_area = row_data.get("photo_area")
        if isinstance(photo_area, dict):
            start = photo_area.get("start")
            end = photo_area.get("end")
            if start and end:
                for m in raw_media:
                    m_coord = m.get("coord", "")
                    if m_coord and is_coord_in_range(m_coord, start, end):
                        return m

        # 3. 座標の隣接性判定
        rep_coord = row_data.get("coordinate_title") or row_data.get("coordinate") or row_data.get("coord")
        if rep_coord:
            r_row, r_col = parse_coord(str(rep_coord))
            if r_row > 0:
                for m in raw_media:
                    m_coord = m.get("coord", "")
                    m_row, m_col = parse_coord(m_coord)
                    if m_col == r_col and abs(m_row - r_row) <= 5:
                        return m
                    if m_row == r_row and abs(m_col - r_col) <= 5:
                        return m

        return None

    def _add_structured_table_with_insights(
        self,
        doc: Document,
        data_list: List[Dict[str, Any]],
        raw_cells: List[Dict[str, Any]],
        raw_media: List[Dict[str, Any]],
        inserted_media: set
    ) -> None:
        """
        テーブルの最右列に「ビジュアルインサイト」テキストカラムを追加して生成する。
        """
        if not data_list or not isinstance(data_list[0], dict):
            for item in data_list:
                doc.add_paragraph(f"・ {item}")
            return

        headers = list(data_list[0].keys())

        # 既存ヘッダーに画像説明系キーがあるかチェックし、なければ「ビジュアルインサイト」列を新規追加
        has_insight_header = any(h in ["visual_insights", "visual_summary", "image_details", "description"] for h in headers)
        new_insight_col_added = False
        if not has_insight_header:
            headers.append("ビジュアルインサイト")
            new_insight_col_added = True

        # Wordテーブルの生成
        table = doc.add_table(rows=len(data_list) + 1, cols=len(headers))
        table.style = 'Table Grid'

        # ヘッダーセルの太字設定
        hdr_cells = table.rows[0].cells
        for idx, header in enumerate(headers):
            hdr_cells[idx].text = str(header)
            for paragraph in hdr_cells[idx].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # データ行の設定
        for r_idx, row_data in enumerate(data_list, 1):
            row_cells = table.rows[r_idx].cells

            # 代表となる行全体の座標を特定 (coordinate_title を最優先)
            row_coord = row_data.get("coordinate_title") or row_data.get("coordinate") or row_data.get("coord")
            if not row_coord:
                for val in row_data.values():
                    if val and not isinstance(val, (dict, list)):
                        if found_coord := self._find_cell_coord(raw_cells, val):
                            row_coord = found_coord
                            break

            # 各データ列に値を書き込み、座標プレフィックスを付与
            for c_idx, header in enumerate(headers):
                if header == "ビジュアルインサイト" and new_insight_col_added:
                    continue

                val = row_data.get(header)
                val_str = str(val) if val is not None else ""

                # 座標プレフィックスの自動生成
                if row_coord:
                    cell_text = f"[{row_coord}] {val_str}"
                else:
                    coord = self._find_cell_coord(raw_cells, val)
                    prefix = f"[{coord}] " if coord else ""
                    cell_text = f"{prefix}{val_str}"

                row_cells[c_idx].text = cell_text

            # ビジュアルインサイトの新規列追加時の書き込み
            if new_insight_col_added:
                insight_cell = row_cells[-1]
                insight_text = ""
                media_item = self._find_associated_media(row_data, raw_media)
                if media_item and media_item.get("visual_summary"):
                    insight_text = str(media_item["visual_summary"])

                # 構造化データ側の既存記述からのフォールバック
                if not insight_text or insight_text == "画像データ":
                    for desc_key in ["image_details", "description", "visual_summary", "details", "visual_insights"]:
                        if row_data.get(desc_key):
                            insight_text = str(row_data[desc_key])
                            break

                insight_cell.text = insight_text if insight_text else "なし"

    def _add_associated_images_below_table(
        self,
        doc: Document,
        sheet_data: Dict[str, Any],
        raw_media: List[Dict[str, Any]],
        inserted_media: set,
        image_width: float
    ) -> None:
        """
        テーブルの直後に、テーブルレコードとの関連を明記したキャプション付きで関連画像を配置する。
        """
        table_data, _ = self._detect_table_data(sheet_data)
        if not table_data:
            return

        # 関連する画像があるか走査
        for r_idx, row_data in enumerate(table_data, 1):
            media_item = self._find_associated_media(row_data, raw_media)
            if not media_item:
                continue

            filename = media_item.get("filename")
            if not filename or filename in inserted_media:
                continue

            image_path = self.media_dir / filename
            if not image_path.exists():
                image_path = self.output_dir / filename

            if image_path.exists():
                # 関連付けを明記したキャプション
                photo_id = row_data.get("id") or row_data.get("photo_no") or r_idx
                clean_id = str(photo_id).replace("No.", "").replace("NO.", "").strip()
                title = row_data.get("title") or row_data.get("name") or "画像データ"

                # 表示座標情報の整理
                coord = row_data.get("coordinate_title") or row_data.get("coordinate") or row_data.get("coord")
                photo_area = row_data.get("photo_area")
                if not coord and isinstance(photo_area, dict):
                    start = photo_area.get("start")
                    end = photo_area.get("end")
                    if start and end:
                        coord = f"{start}-{end}"
                if not coord:
                    coord = "A1"

                m_coord = media_item.get("coord", "")

                caption_text = f"【図：No.{clean_id} {title} (表中座標: {coord} / 画像座標: {m_coord})】"
                doc.add_paragraph(caption_text, style='Caption')

                # 画像挿入
                try:
                    doc.add_picture(str(image_path), width=Inches(image_width))
                    inserted_media.add(filename)
                except Exception as e:
                    logger.error(f"Failed to add associated image {filename}: {e}")
                    doc.add_paragraph(f"[画像挿入エラー: {filename}]")

    def _add_logic_annotations(self, doc: Document, cells: List[Dict[str, Any]]) -> None:
        """
        原本の計算式情報の注釈をQuoteスタイルで追加する。
        """
        RE_LOGIC_FORMULA = re.compile(r"=(SUM|AVERAGE|AVG|COUNT|MAX|MIN|SUBTOTAL|VLOOKUP|IF|ROUND)\b", re.IGNORECASE)

        annotations = []
        for cell in cells:
            formula = cell.get("formula")
            if formula and str(formula).startswith("="):
                if RE_LOGIC_FORMULA.search(str(formula)):
                    coord = cell.get("coord", f"セル({cell['row']},{cell['col']})")
                    unit = cell.get("unit")
                    unit_str = f"（単位: {unit}）" if unit else ""
                    annotations.append(f"ℹ️ セル {coord} は計算式 `{formula}`{unit_str} から導出された集計値です。")

        if annotations:
            for note in annotations:
                doc.add_paragraph(note, style='Quote')

    def _add_image_with_caption(self, doc: Document, media_item: Dict[str, Any], image_width: float) -> None:
        """
        画像をキャプション付きでドキュメントに挿入する (フォールバック用)。
        """
        filename = media_item.get("filename")
        if not filename:
            return

        image_path = self.media_dir / filename
        if not image_path.exists():
            image_path = self.output_dir / filename
            if not image_path.exists():
                return

        coord = media_item.get("coord", "")
        summary = media_item.get("visual_summary", "画像データ")
        if not summary.startswith("【図："):
            summary = summary.replace("[画像概要]", "").strip(": ")
            caption_text = f"【図：{summary} (座標: {coord})】"
        else:
            caption_text = f"{summary[:-1]} (座標: {coord})】"

        doc.add_paragraph(caption_text, style='Caption')

        try:
            doc.add_picture(str(image_path), width=Inches(image_width))
        except Exception as e:
            logger.error(f"Failed to add image {filename}: {e}")
            doc.add_paragraph(f"[画像挿入エラー: {filename}]")
